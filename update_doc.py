# -*- coding: utf-8 -*-
from docx import Document

doc = Document()
doc.add_heading('PyJS Compiler - Informe Léxico, Sintáctico y Semántico', 0)
doc.add_paragraph('Autores: Cristhian Santiago Nastar, Christian Mateo Pastrana')
doc.add_paragraph('Proyecto: Implementación integral (lexer, parser, analizador semántico y UI con Streamlit) para un subconjunto de JavaScript.')
doc.add_paragraph('Fecha de actualización: 09/11/2025')

doc.add_heading('1. Resumen Ejecutivo', level=1)
doc.add_paragraph('El presente documento describe la versión actual del compilador PyJS, el cual procesa un subconjunto de JavaScript mediante tres fases principales (análisis léxico, sintáctico y semántico) y ofrece una interfaz visual en Streamlit para ejecutar pruebas sobre ejemplos incluidos en la carpeta samples. El sistema convierte el código fuente en tokens, construye un Árbol Sintáctico Abstracto (AST), valida el uso correcto de identificadores y tipos, genera una tabla de símbolos y reporta errores con información de línea y columna. Además, se añadió una verificación semántica de división por cero y la validación de miembros permitidos para objetos builtin como console.')

doc.add_heading('2. Arquitectura del Sistema', level=1)
para = doc.add_paragraph()
para.add_run('Componentes principales:\n')
para.add_run('- Analizador Léxico (lexer/lexer.py): utiliza expresiones regulares para clasificar números, cadenas, identificadores, palabras clave, operadores y puntuación.\n')
para.add_run('- Analizador Sintáctico (parser/parser.py): parser de descenso recursivo que reconoce declaraciones de variables, funciones y expresiones aritméticas, construyendo el AST a partir de objetos NodoAST.\n')
para.add_run('- Analizador Semántico (semantic/semantic.py): recorre el AST para construir tablas de símbolos por ámbito, asegura que los identificadores se declaren antes de usarse, valida tipos en operaciones, comprueba miembros válidos de builtins (p. ej. console.log) y detecta divisiones entre cero cuando los operandos son constantes.\n')
para.add_run('- Interfaz de Usuario (UI_Compile.py): aplicativo Streamlit que permite seleccionar/editar ejemplos, compilar y visualizar tokens, AST, tabla de símbolos y errores en pestañas.\n')
para.add_run('- Punto de entrada (main.py): ejecuta las tres fases desde terminal, mostrando tokens, AST y tablas en texto.')

doc.add_heading('3. Lenguaje de Programación Soportado', level=1)
doc.add_paragraph('El subconjunto soportado de JavaScript incluye declaraciones con var/let/const, definiciones de funciones sin parámetros, expresiones aritméticas con +, -, *, /, %, literales numéricos y de cadena, acceso a miembros y llamadas simples. Los tokens reconocidos incluyen palabras clave (var, let, const, function, return, if, else, while, for, true, false, null) e identificadores alfanuméricos estándar.')

doc.add_heading('4. Clases y Módulos Relevantes', level=1)
para = doc.add_paragraph()
para.add_run('Principales clases y responsabilidades:\n')
para.add_run('- Token / Lexer: encapsulan la información léxica y generan la secuencia de tokens desde el código fuente.\n')
para.add_run('- NodoAST / Parser: representan nodos del AST y exponen funciones como parsear(), parsear_declaracion(), parsear_funcion() y parsear_expresion(). Las mejoras recientes permiten tolerar punto y coma opcional antes de una } y reportan identificadores inválidos.\n')
para.add_run('- Symbol / SymbolTable / SemanticAnalyzer: almacenan variables, funciones y builtins por ámbito, detectan duplicados, tipifican expresiones, restringen miembros válidos y agregan la regla de división por cero.\n')
para.add_run('- UI Streamlit: organiza la experiencia en pestañas, colorea el editor, muestra métricas (tokens, errores) y dibuja el AST mediante graphviz integrado sin dependencias externas.')

doc.add_heading('5. Reglas Sintácticas Vigentes', level=1)
doc.add_paragraph('El parser implementa reglas EBNF para Program, FunctionDeclaration, VariableDeclaration, Addition/MulDiv y expresiones primarias (literales, identificadores y paréntesis). Cada regla se asocia a un método específico en parser/parser.py, lo que facilita extender el lenguaje.')

doc.add_heading('6. Análisis Semántico', level=1)
para = doc.add_paragraph()
para.add_run('Las verificaciones actuales incluyen:\n')
para.add_run('- Uso de identificadores declarados y prevención de redefiniciones en el mismo ámbito.\n')
para.add_run('- Consistencia de tipos en operaciones aritméticas: sólo números pueden usarse con -, *, /, %, y la suma prohíbe mezclar number con string salvo concatenación explícita.\n')
para.add_run('- Restricciones sobre símbolos builtin: se registra console con los métodos log, warn y error; referenciar console.hola produce un error.\n')
para.add_run('- Detección de división por cero cuando el divisor se evalúa como constante.\n')
para.add_run('- Generación de una tabla de símbolos por scope (global, funciones, bloques).')

doc.add_heading('7. Interfaz Visual (UI_Compile.py)', level=1)
doc.add_paragraph('La UI en Streamlit incorpora la paleta Py/JS para el encabezado, un editor con estilo tipo IDE, métricas rápidas y pestañas: Tokens y símbolos (tablas organizadas), AST (gráfico generado con graphviz_chart) y Errores (separados por sintácticos/semánticos). Además, permite cargar ejemplos desde samples, editarlos, guardar cambios opcionalmente y lanzar el compilador con un botón único.')

doc.add_heading('8. Ejemplos de Prueba', level=1)
para = doc.add_paragraph()
para.add_run('Cada archivo en samples incluye comentarios que indican cómo generar errores sintácticos y semánticos. Ejemplos clave:\n')
para.add_run('- ejemplo.js: cambiar console.log por console.hola para detonar la validación de miembros.\n')
para.add_run('- multiplicacion.js: reemplazar x * y por x * "4" provoca un error de tipos.\n')
para.add_run('- promedio.js: sustituir el divisor por 0 activa la detección de división por cero.\n')
para.add_run('- suma.js: definir una variable con identificador vacío demuestra la detección sintáctica y evita que la tabla de símbolos registre entradas inválidas.')

doc.add_heading('9. Código Fuente y Ejecución', level=1)
doc.add_paragraph('El repositorio mantiene los módulos main.py, lexer/lexer.py, parser/parser.py, semantic/semantic.py, UI_Compile.py y streamlit_app.py. Para ejecutar desde consola se usa python main.py; para la UI se ejecuta streamlit run UI_Compile.py. Ambos flujos utilizan la misma lógica de compilación, garantizando resultados consistentes.')

doc.add_heading('10. Conclusiones', level=1)
doc.add_paragraph('El proyecto evolucionó de un analizador léxico/sintáctico básico a un compilador educativo completo con soporte semántico y representación visual. Las mejoras claves incluyen la tabla de símbolos jerárquica, la validación de miembros builtin, la detección temprana de división por cero y la interfaz Streamlit que agiliza las pruebas de estudiantes y docentes. El documento queda alineado con el estado actual del repositorio al 09/11/2025 y servirá como base para futuras extensiones (manejo de parámetros, retorno de funciones o generación de código intermedio).')

doc.save('Informe_Analizador_Lexico_Sintactico.docx')
